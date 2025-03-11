from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


def generate_docx_report(case_id, forensic_data):
    doc = Document()
    doc.add_heading('Forensic Investigation Report', level=1)

    # Executive Summary
    case_info = forensic_data["case"]
    doc.add_heading('1. Executive Summary', level=2)
    doc.add_paragraph(f"Case ID: {case_info['case_id']}")
    doc.add_paragraph(f"Investigator: {case_info['investigator']}")
    doc.add_paragraph(f"Summary: {case_info['summary']}")
    doc.add_paragraph(f"Start Time: {case_info['start_time']}")
    doc.add_paragraph(f"End Time: {case_info.get('end_time', 'Ongoing')}")

    # Memory Analysis
    doc.add_heading('2. Memory Analysis Findings', level=2)
    for entry in forensic_data["ram"]:
        doc.add_paragraph(f"Process: {entry['process_name']} | PID: {entry['pid']} | Suspicious: {entry['suspicious']}")

    # File System Findings
    doc.add_heading('3. File System Findings', level=2)
    for entry in forensic_data["files"]:
        doc.add_paragraph(f"File: {entry['file_name']} | Path: {entry['file_path']} | Suspicious: {entry['suspicious']}")

    # Threat Intelligence
    doc.add_heading('4. Threat Intelligence & Risk Levels', level=2)
    for entry in forensic_data["threats"]:
        doc.add_paragraph(f"Threat: {entry['threat_type']} | Risk Level: {entry['risk_level']} | Detected: {entry['detected_at']}")

    # Save the report
    report_name = f"forensic_report_case_{case_id}.docx"
    doc.save(report_name)
    return report_name

def generate_pdf_report(case_id, forensic_data):
    pdf_name = f"forensic_report_case_{case_id}.pdf"
    c = canvas.Canvas(pdf_name, pagesize=letter)
    c.drawString(100, 750, f"Forensic Investigation Report - Case {case_id}")

    # Executive Summary
    y_position = 730
    c.drawString(100, y_position, f"Investigator: {forensic_data['case']['investigator']}")
    y_position -= 20
    c.drawString(100, y_position, f"Summary: {forensic_data['case']['summary']}")
    
    # Memory Analysis
    y_position -= 40
    c.drawString(100, y_position, "Memory Analysis Findings:")
    for entry in forensic_data["ram"]:
        y_position -= 20
        c.drawString(100, y_position, f"Process: {entry['process_name']} | Suspicious: {entry['suspicious']}")

    # Save the report
    c.save()
    return pdf_name

def generate_md_report(case_id, forensic_data):
    md_name = f"forensic_report_case_{case_id}.md"
    with open(md_name, "w") as f:
        f.write(f"# Forensic Investigation Report - Case {case_id}\n\n")
        f.write(f"## 1. Executive Summary\n")
        f.write(f"**Investigator:** {forensic_data['case']['investigator']}\n")
        f.write(f"**Summary:** {forensic_data['case']['summary']}\n")
        
        f.write(f"## 2. Memory Analysis Findings\n")
        for entry in forensic_data["ram"]:
            f.write(f"- Process: {entry['process_name']} | Suspicious: {entry['suspicious']}\n")

    return md_name
